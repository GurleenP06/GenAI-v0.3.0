import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Dict

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA_DIR = Path("./chat_data")


def export_as_txt(chat_history: List[Dict]) -> Path:
    content = ""
    for msg in chat_history:
        content += f"{msg['role'].upper()}: {msg['message']}\n\n"

    temp_path = DATA_DIR / f"export_{uuid.uuid4()}.txt"
    with open(temp_path, "w", encoding="utf-8") as f:
        f.write(content)
    return temp_path


def markdown_to_docx(chat_history: List[Dict]) -> Path:
    doc = DocxDocument()

    title = doc.add_heading('Chat Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    for msg in chat_history:
        role = msg['role'].capitalize()
        content = msg['message']

        p = doc.add_paragraph()
        run = p.add_run(f"{role}: ")
        run.bold = True

        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip():
                doc.add_paragraph(line)

        doc.add_paragraph("")

    temp_path = DATA_DIR / f"export_{uuid.uuid4()}.docx"
    doc.save(temp_path)
    return temp_path
